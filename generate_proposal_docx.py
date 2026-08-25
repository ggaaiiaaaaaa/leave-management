import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_proposal():
    doc = docx.Document()

    # Set 1-inch margins (72pt)
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    # Base font styling: Arial or Calibri or Segoe UI
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(30, 41, 59) # Slate 800

    # ---------------------------------------------------------
    # TITLE & HEADER
    # ---------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run("JTYEO CPA ACCOUNTING OFFICE")
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(15, 39, 68) # Deep Executive Navy

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(2)
    run_sub = p_sub.add_run("Leave & Attendance Management System — Feature Overview & Pricing Proposal")
    run_sub.font.size = Pt(11)
    run_sub.font.color.rgb = RGBColor(71, 85, 105)

    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after = Pt(18)
    run_date = p_date.add_run("Prepared August 2026")
    run_date.font.size = Pt(9)
    run_date.font.color.rgb = RGBColor(148, 163, 184)

    # ---------------------------------------------------------
    # 1. FEATURE REQUIREMENTS
    # ---------------------------------------------------------
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    r_h1 = h1.add_run("1. Feature Requirements")
    r_h1.font.size = Pt(14)
    r_h1.font.bold = True
    r_h1.font.color.rgb = RGBColor(15, 39, 68)

    # --- Role 1: Managing Partner & HR Admin ---
    p_admin_head = doc.add_paragraph()
    p_admin_head.paragraph_format.space_before = Pt(8)
    p_admin_head.paragraph_format.space_after = Pt(3)
    r_ah = p_admin_head.add_run("Managing Partner & HR Admin")
    r_ah.font.size = Pt(12)
    r_ah.font.bold = True
    r_ah.font.color.rgb = RGBColor(15, 39, 68)

    admin_features = [
        "Executive workforce dashboard with live headcount, on-duty percentage, and active leaves overview",
        "Firm-wide leave ledger tracking all filed, approved, rejected, and active staff absences",
        "Statutory DOLE compliance tracker (Article 95 Service Incentive Leave 5-day monetization reserve)",
        "Review and decision engine for all pending leave applications across departments",
        "Approver feedback dialog allowing personalized notes or explanations on decisions",
        "Manual balance adjustment tool to increment (+) or decrement (-) credits with mandatory audit reasoning",
        "Global Monthly Accrual engine (+1.25 Vacation Leave days per month for active staff)",
        "Annual DOLE SIL Rollover & Reset engine refreshing staff to 5.0 statutory days at fiscal year-end",
        "One-click DOLE/BIR-ready Payroll CSV export with employee IDs, deduction codes, and paid/unpaid status",
        "Interactive monthly team leave calendar with color-coded statutory categories and Philippine public holidays",
        "Immutable system audit trail logging leave submissions, approvals/rejections, and balance adjustments with IP timestamps",
        "Engagement team coverage matrix across Audit & Assurance, Taxation & Compliance, and Bookkeeping"
    ]
    for feat in admin_features:
        p = doc.add_paragraph(feat, style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2.5)

    # --- Role 2: Senior Lead / Engagement Supervisor ---
    p_sup_head = doc.add_paragraph()
    p_sup_head.paragraph_format.space_before = Pt(10)
    p_sup_head.paragraph_format.space_after = Pt(3)
    r_sh = p_sup_head.add_run("Senior Lead / Engagement Supervisor")
    r_sh.font.size = Pt(12)
    r_sh.font.bold = True
    r_sh.font.color.rgb = RGBColor(15, 39, 68)

    sup_features = [
        "Dedicated lead reviewer approvals queue displaying pending staff requests front-and-center",
        "Live conflict and overlap detection to prevent understaffing during critical client audit engagements",
        "One-click Approval / Rejection modal with reviewer note support (e.g., coverage handover instructions)",
        "Personal leave balance monitor (DOLE SIL, Vacation Leave, Sick Leave, Solo Parent)",
        "Submission of personal leave requests with automated working day calculation",
        "Department-level team roster and attendance matrix view",
        "Interactive full calendar visibility into firm-wide schedule and DOLE holidays"
    ]
    for feat in sup_features:
        p = doc.add_paragraph(feat, style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2.5)

    # --- Role 3: Staff Accountant / Associate CPA ---
    p_staff_head = doc.add_paragraph()
    p_staff_head.paragraph_format.space_before = Pt(10)
    p_staff_head.paragraph_format.space_after = Pt(3)
    r_sth = p_staff_head.add_run("Staff Accountant / Associate CPA")
    r_sth.font.size = Pt(12)
    r_sth.font.bold = True
    r_sth.font.color.rgb = RGBColor(15, 39, 68)

    staff_features = [
        "Self-service dashboard displaying real-time personal allowances (SIL, Vacation, Sick, Solo Parent, Magna Carta)",
        "Leave application form with smart working-day calculator excluding weekends and Philippine public holidays",
        "Flexible duration modes: Full Day (1.0d), Half-Day Morning (0.5d), Half-Day Afternoon (0.5d)",
        "Medical certificate & supporting slip file uploader (required for sick leaves > 2 consecutive working days)",
        "Personal leave application history with live status tracking (Pending, Approved, Rejected) and signoff notes",
        "Firm leave & holiday calendar access to check team availability before filing",
        "Comprehensive DOLE statutory leave rules and legal entitlement reference guide"
    ]
    for feat in staff_features:
        p = doc.add_paragraph(feat, style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2.5)

    # ---------------------------------------------------------
    # 2. PRICING BREAKDOWN
    # ---------------------------------------------------------
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(16)
    h2.paragraph_format.space_after = Pt(6)
    r_h2 = h2.add_run("2. Updated Pricing Breakdown")
    r_h2.font.size = Pt(14)
    r_h2.font.bold = True
    r_h2.font.color.rgb = RGBColor(15, 39, 68)

    p_pricing_desc = doc.add_paragraph()
    p_pricing_desc.paragraph_format.space_before = Pt(0)
    p_pricing_desc.paragraph_format.space_after = Pt(10)
    p_pricing_desc.add_run(
        "The leave and attendance workflows for a CPA practice require strict alignment with Philippine Labor Standards (DOLE Art. 95) "
        "and client engagement deadlines. The Tier 1 base system provides complete paperless leave governance and statutory compliance. "
        "Tier 2 (₱70,000) represents the Recommended Sweet Spot, introducing automated 2-tier partner signoffs, 1-click email approvals, "
        "and peak BIR Tax Season lockout rules. Tier 3 (₱85,000) offers enterprise-grade payroll system sync, morning Viber/Teams emergency alerts, "
        "and 3 months priority SLA maintenance."
    )

    # Create Comparison Table: 3 rows x 4 cols
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Widths: Col 0 = 1.0 in, Col 1 = 1.8 in, Col 2 = 1.85 in, Col 3 = 1.85 in (Total = 6.5 in)
    col_widths = [Inches(0.9), Inches(1.8), Inches(1.9), Inches(1.9)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # Row 0: Headers
    r0 = table.rows[0]
    r0.cells[0].text = ""
    r0.cells[1].text = "TIER 1\nCore System"
    r0.cells[2].text = "TIER 2 (Sweet Spot)\nPro Practice Suite"
    r0.cells[3].text = "TIER 3\nEnterprise Suite"

    for idx, cell in enumerate(r0.cells):
        shading = parse_xml(r'<w:shd {} w:fill="0F2744"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(9.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    # Row 1: Prices
    r1 = table.rows[1]
    r1.cells[0].text = "Price"
    r1.cells[1].text = "₱50,000"
    r1.cells[2].text = "₱70,000"
    r1.cells[3].text = "₱85,000"

    for idx, cell in enumerate(r1.cells):
        shading = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.bold = True
                r.font.color.rgb = RGBColor(15, 39, 68) if idx == 0 else (RGBColor(2, 132, 199) if idx == 2 else RGBColor(15, 23, 42))

    # Row 2: Scope Items
    r2 = table.rows[2]
    r2.cells[0].text = "Scope"
    
    t1_text = (
        "• Centralized leave & attendance database\n"
        "• DOLE statutory leave tracking (SIL 5d, VL 12d, SL 10d, Solo Parent 7d)\n"
        "• Role-separated dashboards (Staff CPA, Senior Lead, Partner/HR)\n"
        "• Philippine holiday & weekend auto-exclusion working day calculator\n"
        "• Interactive leave calendar (FullCalendar.js)\n"
        "• Managerial decision modal with feedback note support\n"
        "• Basic CSV leave ledger export\n"
        "• Targeted audit trail (applications, decisions, adjustments)\n"
        "• 1-click demo persona switcher"
    )
    r2.cells[1].text = t1_text

    t2_text = (
        "• Everything included in Tier 1\n"
        "• 2-Tier Approval Workflow (Senior Audit Lead -> Managing Partner)\n"
        "• Automated Email Notifications with 1-click mobile approve/reject links\n"
        "• Peak BIR Tax Season 'Leave Freeze' (Automated Blackout Period controls)\n"
        "• DOLE SIL Monetization & Cash Conversion Calculator\n"
        "• Medical Certificate & Supporting Slip Cloud Storage & in-app preview\n"
        "• HR Leave Credit Adjuster (+/- days) with administrative audit reasons\n"
        "• Global Monthly Accrual engine (+1.25d VL/month)\n"
        "• Annual DOLE SIL Reset (5.0d) fiscal year rollover"
    )
    r2.cells[2].text = t2_text

    t3_text = (
        "• Everything included in Tier 2\n"
        "• 1-Click Payroll & Timekeeper Sync Engine (mapped to payroll Excel/software)\n"
        "• Department Engagement Matrix with minimum staffing thresholds\n"
        "• Instant Chatbot / Messaging Alerts (Viber / MS Teams / Slack for morning absences)\n"
        "• Partner Executive BI Analytics (Absentee heatmaps, seasonal utilization rates)\n"
        "• Full source code handover with cloud/server deployment setup\n"
        "• 3 Months Priority SLA Support & warranty\n"
        "• Complete Video Training Walkthrough library for staff onboarding"
    )
    r2.cells[3].text = t3_text

    for idx, cell in enumerate(r2.cells):
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            for r in p.runs:
                if idx == 0:
                    r.font.size = Pt(9.5)
                    r.font.bold = True
                else:
                    r.font.size = Pt(8.5)
                    r.font.color.rgb = RGBColor(51, 65, 85)

    # Style Table Borders (Subtle Gray Grid)
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        r'<w:tblBorders {} >'
        r'<w:top w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="CBD5E1"/>'
        r'<w:left w:val="none"/>'
        r'<w:right w:val="none"/>'
        r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        r'</w:tblBorders>'.format(nsdecls('w'))
    )
    tblPr.append(tblBorders)

    # ---------------------------------------------------------
    # 3. TERMS & CONDITIONS
    # ---------------------------------------------------------
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(16)
    h3.paragraph_format.space_after = Pt(6)
    r_h3 = h3.add_run("3. Terms & Conditions")
    r_h3.font.size = Pt(14)
    r_h3.font.bold = True
    r_h3.font.color.rgb = RGBColor(15, 39, 68)

    terms_data = [
        ("Payment Schedule", [
            "A 50% down payment is required to commence project development, with the remaining 50% due upon final system deployment and client acceptance.",
            "Tier upgrades or additional custom add-on modules requested after project kickoff will be quoted and invoiced separately upon agreement.",
            "All prices are quoted in Philippine Pesos (₱) and are net of applicable withholding taxes unless otherwise stated."
        ]),
        ("Project Timeline & Delivery", [
            "Tier 1 (Core Portal): Estimated delivery within two to three (2–3) weeks from kickoff.",
            "Tier 2 (Pro Practice Suite): Estimated delivery within three to four (3–4) weeks from kickoff.",
            "Tier 3 (Enterprise Suite): Estimated delivery within four to five (4–5) weeks from kickoff.",
            "Timelines assume timely provision of firm roster details, engagement policies, and milestone signoffs."
        ]),
        ("Revisions & Milestone Signoffs", [
            "Two (2) formal rounds of design and workflow revisions are included per development milestone.",
            "Additional revisions or scope changes outside the approved functional specifications will be quoted separately."
        ]),
        ("Ownership & Confidentiality", [
            "Full ownership of the completed application, including custom source code and database schemas, is transferred to JTYeo CPA Accounting Office upon receipt of final payment.",
            "All staff records, client engagement lists, partner communications, and firm financial data are treated under strict non-disclosure."
        ]),
        ("Support, SLA & Maintenance", [
            "A complimentary 30-day post-launch warranty is included for Tier 1 and Tier 2 for bug fixes and stability assurance.",
            "Tier 3 includes an extended three (3) months of Priority SLA Support covering feature fine-tuning, partner support, and staff onboarding.",
            "Annual maintenance, server hosting management, and continuous feature updates beyond warranty are available under an optional retainer agreement."
        ]),
        ("Validity of Proposal", [
            "This proposal, including all presented tier pricing, is valid for thirty (30) days from the date of presentation."
        ])
    ]

    for section_title, bullets in terms_data:
        p_sec = doc.add_paragraph()
        p_sec.paragraph_format.space_before = Pt(8)
        p_sec.paragraph_format.space_after = Pt(2)
        r_sec = p_sec.add_run(section_title)
        r_sec.font.size = Pt(11)
        r_sec.font.bold = True
        r_sec.font.color.rgb = RGBColor(15, 39, 68)

        for b in bullets:
            p_b = doc.add_paragraph(b, style='List Paragraph')
            p_b.paragraph_format.space_before = Pt(1)
            p_b.paragraph_format.space_after = Pt(2.5)

    doc.save("JTYeo_CPA_Leave_Management_Proposal.docx")
    print("Successfully generated JTYeo_CPA_Leave_Management_Proposal.docx!")

if __name__ == "__main__":
    create_proposal()
